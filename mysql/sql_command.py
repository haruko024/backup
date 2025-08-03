from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Set the title
title = doc.add_heading('SQL Commands List', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Define SQL command categories and their commands
sql_commands = {
    "Data Definition Language (DDL)": [
        "CREATE DATABASE – Creates a new database.",
        "CREATE TABLE – Creates a new table.",
        "ALTER TABLE – Modifies an existing table.",
        "DROP TABLE – Deletes a table.",
        "DROP DATABASE – Deletes a database.",
        "TRUNCATE TABLE – Removes all data from a table."
    ],
    "Data Manipulation Language (DML)": [
        "SELECT – Retrieves data from a table.",
        "INSERT INTO – Adds new records.",
        "UPDATE – Modifies existing records.",
        "DELETE – Deletes records."
    ],
    "Data Control Language (DCL)": [
        "GRANT – Gives user access privileges.",
        "REVOKE – Removes user access privileges."
    ],
    "Transaction Control Language (TCL)": [
        "BEGIN TRANSACTION – Starts a transaction.",
        "COMMIT – Saves all changes.",
        "ROLLBACK – Undoes changes since last commit.",
        "SAVEPOINT – Sets a point to roll back to.",
        "SET TRANSACTION – Defines transaction properties."
    ],
    "Clauses and Operators": [
        "WHERE – Filters records.",
        "ORDER BY – Sorts the result.",
        "GROUP BY – Groups rows.",
        "HAVING – Filters groups.",
        "JOIN – Combines rows from tables (INNER, LEFT, RIGHT, FULL).",
        "UNION / UNION ALL – Combines results of two queries.",
        "LIMIT / OFFSET – Restricts number of results.",
        "IN, BETWEEN, LIKE, IS NULL – Conditional operators."
    ],
    "Functions": [
        "COUNT(), SUM(), AVG(), MIN(), MAX() – Aggregate functions.",
        "UPPER(), LOWER(), CONCAT(), SUBSTRING() – String functions.",
        "NOW(), CURDATE(), DATEDIFF() – Date/time functions.",
        "COALESCE(), NULLIF(), CAST() – Other useful functions."
    ]
}

# Add content to the document
for category, commands in sql_commands.items():
    doc.add_heading(category, level=1)
    for command in commands:
        p = doc.add_paragraph(command, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)

# Save the document
doc.save("SQL_Commands_List.docx")

print("Document created: SQL_Commands_List.docx")
