import os
import subprocess
import requests
import speech_recognition as sr
from fpdf import FPDF
from datetime import datetime
from docx import Document
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from colorama import Fore, init
import ast
import shutil

init(autoreset=True)

# === CONFIGURATION ===
API_KEY = "sk_e4f3e667710caaec91adc68f2d8bbb77f4045b5cfe471090"
VOICE_ID = "EXAVITQu4vr4xnSDxMaL"

# === TEXT TO SPEECH ===
def speak(text):
    url = f"https://api.elevenlabs.io/v1/text-to-speech/{VOICE_ID}"
    headers = {
        "xi-api-key": API_KEY,
        "Content-Type": "application/json"
    }
    data = {
        "text": text,
        "model_id": "eleven_monolingual_v1",
        "voice_settings": {
            "stability": 0.75,
            "similarity_boost": 0.75
        }
    }

    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        with open("output.mp3", "wb") as f:
            f.write(response.content)
        try:
            os.system("mpg123 output.mp3")
        except Exception as e:
            print("[AUDIO ERROR]", e)
    else:
        print("[TTS ERROR]", response.status_code, response.json())

# === LISTEN FUNCTION ===
def listen():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print(Fore.YELLOW + "[PaulAI] Listening...")
        audio = recognizer.listen(source)
    try:
        command = recognizer.recognize_google(audio)
        print(Fore.GREEN + f"[You] >>> {command}")
        return command.lower().strip()
    except sr.UnknownValueError:
        speak("Sorry, I didn't catch that. Please repeat.")
        return listen()
    except sr.RequestError:
        speak("Speech recognition service is not available.")
        return ""

# === BASIC DEMO DOCX/PDF ===
def create_docx(content="Hello from PaulAI! This is your DOCX file."):
    doc = Document()
    doc.add_heading("PaulAI Document", level=1)
    doc.add_paragraph(content)
    doc.save("output.docx")
    print(Fore.CYAN + "[✓] DOCX file created.")

def convert_to_pdf(text="This is your PDF content."):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(190, 10, text)
    pdf.output("output.pdf")
    print(Fore.CYAN + "[✓] PDF file created.")

# === ADVANCED AUTOMATION ===
def something():
    speak("Say Something?")
    command = listen()

    if "kupal" in command:
        speak("Kupal! Wait for a while...")
    else:
        speak("You didn't say kupal, but I'm doing it anyway...")

    docx_and_pdf()
    merge_docx_pdf()
    speak("All tasks completed successfully.")
    speak("Thanks for running this code. Please subscribe to my GitHub and visit my website.")
    print(Fore.GREEN + "GitHub : https://github.com/paulmendoza24")
    print(Fore.GREEN + "Website : https://paulmendoza24.github.io/index/")
    speak("Remember!: With great power! comes great responsibility.")

# === LOAD CONFIGURATION ===
date = datetime.now().strftime("%B %d, %Y")

with open("names.txt", "r", encoding="utf-8") as f:
    names = [line.strip() for line in f if line.strip()]

with open("config.txt", "r", encoding="utf-8") as f:
    lines = f.read().split("@")
    temp = lines[0].split("=")[1].strip()
    cont = lines[1].split("=")[1].strip()

# === DOCX + PDF GENERATION ===
def docxonly():
    speak(f"{len(names)} name list found! Please wait...")
    speak("Generating DOCX files.")
    for index, name in enumerate(names):
        doc = DocxTemplate(f"templates/{temp}")
        try:
            context = ast.literal_eval("{" + cont + "}")
        except Exception as e:
            print(Fore.RED + "[Context Error]", e)
            return

        filename = f"{name}_output.docx" if index == 0 else f"{name}_output{index}.docx"
        os.makedirs("output/docx", exist_ok=True)
        doc.render(context)
        doc.save(f"output/docx/{filename}")
        print(Fore.CYAN + f"⭕ Generated {filename}")
    speak("DOCX files created.")

def generatepdf():
    if not shutil.which("libreoffice"):
        speak("LibreOffice is not installed. Cannot convert to PDF.")
        return

    speak("Converting DOCX files to PDF.")
    docx_folder = os.path.abspath("output/docx")
    pdf_folder = os.path.abspath("output/pdfs")
    os.makedirs(pdf_folder, exist_ok=True)

    docx_files = [f for f in os.listdir(docx_folder) if f.endswith(".docx")]
    if not docx_files:
        speak("No DOCX files found.")
        return

    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", pdf_folder
    ] + [os.path.join(docx_folder, f) for f in docx_files],
    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    speak("PDF conversion completed.")

def docx_and_pdf():
    speak("Generating DOCX and PDF.")
    docxonly()
    generatepdf()

# === MERGING ===
def merge_docx():
    speak("Merging DOCX files.")
    docx_folder = "output/docx"
    docx_files = sorted([
        os.path.join(docx_folder, f)
        for f in os.listdir(docx_folder)
        if f.endswith(".docx") and f != "merged.docx"
    ])

    if not docx_files:
        speak("No DOCX files found to merge.")
        return

    master = Document(docx_files[0])
    composer = Composer(master)
    for file in docx_files[1:]:
        composer.append(Document(file))

    os.makedirs("output/merged", exist_ok=True)
    merged_path = os.path.join("output/merged", "merged.docx")
    composer.save(merged_path)
    print(Fore.CYAN + "⭕ Generated merged.docx")
    speak("DOCX files merged successfully.")

def merge_docx_pdf():
    speak("Merging DOCX and converting to PDF.")
    merge_docx()
    merged_docx_path = os.path.abspath("output/merged/merged.docx")
    merged_pdf_output = os.path.abspath("output/merged")

    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf", "--outdir", merged_pdf_output,
        merged_docx_path
    ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    print(Fore.CYAN + "⭕ Generated merged.pdf")
    speak("Merged DOCX converted to PDF.")

# === MAIN MENU ===
def menu():
    speak("Hello! I am your document assistant. What would you like me to do, sir Paul?")
    print(Fore.YELLOW + "[PaulAI] You can say: generate docx, convert to PDF, generate both, full automation, help, or exit.")

    while True:
        command = listen()
        if "generate docx" in command:
            create_docx()
            speak("Your DOCX file has been created.")
        elif "convert to pdf" in command:
            convert_to_pdf()
            speak("Your PDF file has been created.")
        elif "generate both" in command:
            create_docx()
            convert_to_pdf()
            speak("Both DOCX and PDF files have been generated.")
        elif "full automation" in command or "complete" in command:
            something()
            break
        elif "help" in command or "commands" in command:
            print(Fore.YELLOW + "[PaulAI] You can say: generate docx, convert to PDF, generate both, full automation, or exit.")
            speak("You can say: generate docx, convert to PDF, generate both, full automation, or exit.")
        elif "exit" in command:
            speak("Goodbye, sir Paul!")
            break
        else:
            speak("Sorry, I did not understand that. Please try again.")

# === ENTRY POINT ===
if __name__ == "__main__":
    menu()
